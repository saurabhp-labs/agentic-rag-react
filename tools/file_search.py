"""
File system search tool - grep and read files directly from the data directory.
This is the "file-first" approach: no chunking/embedding loss, agent sees raw content.
"""
import os
import re
from llama_index.core.tools import FunctionTool

from config import config


def search_files(query: str) -> str:
    """
    Search through files in the data directory for lines matching the query.
    Searches case-insensitively across all text-readable files (.txt, .docx, .md).
    Returns matching lines with their file name and line number.

    Args:
        query: The search term or phrase to look for in files.
    """
    data_dir = config.data_dir
    if not os.path.exists(data_dir):
        return f"Data directory not found: {data_dir}"

    results = []
    pattern = re.compile(re.escape(query), re.IGNORECASE)

    for filename in os.listdir(data_dir):
        file_path = os.path.join(data_dir, filename)
        if not os.path.isfile(file_path):
            continue

        ext = os.path.splitext(filename)[1].lower()

        # Read text-based files
        content = None
        if ext == ".txt":
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
            except UnicodeDecodeError:
                with open(file_path, "r", encoding="latin-1") as f:
                    content = f.read()
        elif ext == ".docx":
            try:
                import docx
                doc = docx.Document(file_path)
                content = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            except Exception:
                continue

        if content is None:
            continue

        # Search line by line
        for i, line in enumerate(content.split("\n"), 1):
            if pattern.search(line):
                results.append(f"[{filename}:L{i}] {line.strip()}")

    if not results:
        return f"No matches found for '{query}' in any files."

    # Cap results to avoid overwhelming the LLM
    if len(results) > 20:
        return "\n".join(results[:20]) + f"\n\n... and {len(results) - 20} more matches."
    return "\n".join(results)


def read_file(filename: str) -> str:
    """
    Read the full content of a specific file from the data directory.
    Use this after search_files to read the complete context of a matching file.

    Args:
        filename: The name of the file to read (e.g. '05_EV_Market_Research_Report.txt').
    """
    file_path = os.path.join(config.data_dir, filename)

    if not os.path.exists(file_path):
        available = [f for f in os.listdir(config.data_dir) if os.path.isfile(os.path.join(config.data_dir, f))]
        return f"File '{filename}' not found. Available files:\n" + "\n".join(f"  - {f}" for f in available)

    ext = os.path.splitext(filename)[1].lower()

    if ext == ".txt":
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="latin-1") as f:
                content = f.read()
    elif ext == ".docx":
        try:
            import docx
            doc = docx.Document(file_path)
            content = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            return f"Error reading {filename}: {e}"
    else:
        return f"Cannot read file type '{ext}'. Supported: .txt, .docx"

    # Truncate very long files
    if len(content) > 8000:
        return content[:8000] + f"\n\n... [truncated, {len(content)} total characters]"
    return content


def create_file_search_tools() -> list:
    """Create file search tools for the ReAct agent."""
    return [
        FunctionTool.from_defaults(
            fn=search_files,
            name="file_search",
            description=(
                "Search through files in the data directory using keyword/phrase matching. "
                "Use this for exact term lookups, finding specific names, numbers, or phrases. "
                "Better than vector_search when you know the exact term to look for."
            ),
        ),
        FunctionTool.from_defaults(
            fn=read_file,
            name="file_read",
            description=(
                "Read the full content of a specific file from the data directory. "
                "Use this after file_search to get complete context, or when you know "
                "which file contains the answer. Input is the filename."
            ),
        ),
    ]
